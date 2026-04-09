"""
EnterpriseCore — Générateur de rapports Word (.docx)
=====================================================
"""

import io
import json
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Couleurs ──────────────────────────────────────────────────
COLOR_ACCENT  = RGBColor(0x5B, 0x6A, 0xF5)   # #5b6af5
COLOR_HIGH    = RGBColor(0xEF, 0x44, 0x44)   # rouge
COLOR_MEDIUM  = RGBColor(0xF5, 0x9E, 0x0B)   # orange
COLOR_LOW     = RGBColor(0x22, 0xC5, 0x5E)   # vert
COLOR_UNKNOWN = RGBColor(0x64, 0x74, 0x8B)   # gris
COLOR_DARK    = RGBColor(0x1E, 0x29, 0x3B)   # titre foncé
COLOR_MUTED   = RGBColor(0x64, 0x74, 0x8B)   # texte muted


def _risk_color(level: str) -> RGBColor:
    return {
        "HIGH":   COLOR_HIGH,
        "MEDIUM": COLOR_MEDIUM,
        "LOW":    COLOR_LOW,
    }.get(level, COLOR_UNKNOWN)


def _set_cell_bg(cell, hex_color: str):
    """Définit la couleur de fond d'une cellule."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = COLOR_ACCENT if level == 1 else COLOR_DARK
    run.font.bold = True


def _add_paragraph(doc: Document, text: str, muted: bool = False, italic: bool = False) -> None:
    p   = doc.add_paragraph(text)
    run = p.runs[0] if p.runs else p.add_run(text)
    if muted:
        run.font.color.rgb = COLOR_MUTED
        run.font.size      = Pt(10)
    if italic:
        run.font.italic = True


def generate_mission_docx(mission: dict, tasks: list) -> bytes:
    """
    Génère un document Word à partir des données de mission.
    Retourne les bytes du fichier .docx.
    """
    doc = Document()

    # ── Marges ───────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # ── Page de titre ─────────────────────────────────────────
    risk       = mission.get("final_risk_level", "UNKNOWN")
    risk_color = _risk_color(risk)
    date_str   = str(mission.get("created_at", ""))[:10]

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("RAPPORT D'ANALYSE JURIDIQUE")
    run.bold           = True
    run.font.size      = Pt(22)
    run.font.color.rgb = COLOR_ACCENT

    doc.add_paragraph()

    mission_title_para = doc.add_paragraph()
    mission_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = mission_title_para.add_run(mission.get("title", "Mission sans titre"))
    run2.bold           = True
    run2.font.size      = Pt(16)
    run2.font.color.rgb = COLOR_DARK

    doc.add_paragraph()

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f"Mission #{mission.get('id', '')}  ·  {date_str}  ·  Risque global : {risk}"
    )
    meta_run.font.color.rgb = risk_color
    meta_run.font.size      = Pt(12)
    meta_run.bold           = True

    doc.add_page_break()

    # ── Résumé exécutif ───────────────────────────────────────
    _add_heading(doc, "1. Résumé exécutif", level=1)

    report = mission.get("final_report") or {}
    if isinstance(report, str):
        try:
            report = json.loads(report)
        except Exception:
            report = {}

    exec_summary = report.get("executive_summary", "")
    if exec_summary:
        _add_paragraph(doc, exec_summary)
    else:
        _add_paragraph(doc, mission.get("objective", ""), muted=True)

    doc.add_paragraph()

    # Matrice de risque (tableau)
    _add_heading(doc, "Indicateurs clés", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, label in enumerate(["Mission", "Risque", "Tâches complétées", "Confiance"]):
        hdr_cells[i].text = label
        _set_cell_bg(hdr_cells[i], "5B6AF5")
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True

    row_cells = table.add_row().cells
    conf = report.get("average_confidence", 0)
    row_cells[0].text = mission.get("title", "")[:40]
    row_cells[1].text = risk
    row_cells[2].text = f"{mission.get('completed_tasks', 0)} / {mission.get('tasks_count', 0)}"
    row_cells[3].text = f"{int(float(conf) * 100)}%" if conf else "—"

    doc.add_paragraph()

    # Actions prioritaires
    key_actions = report.get("key_actions", [])
    if key_actions:
        _add_heading(doc, "Actions prioritaires", level=2)
        for i, action in enumerate(key_actions, 1):
            p   = doc.add_paragraph(style="List Number")
            run = p.add_run(action)
            run.font.size = Pt(11)

    doc.add_page_break()

    # ── Analyses par agent ────────────────────────────────────
    _add_heading(doc, "2. Analyses détaillées par agent", level=1)

    for task in tasks:
        if task.get("status") != "done":
            continue

        result = task.get("result") or {}
        if isinstance(result, str):
            try:
                result = json.loads(result)
            except Exception:
                result = {}

        task_risk   = task.get("risk_level", "UNKNOWN")
        task_conf   = task.get("confidence", 0)
        agent_type  = task.get("agent_type", "").upper()

        _add_heading(doc, f"{task.get('task_order', '')}.  {task.get('task_title', '')}", level=2)

        # Méta-info
        meta = doc.add_paragraph()
        meta.add_run(f"Agent : {agent_type}  ·  ").bold = False
        run_risk = meta.add_run(f"Risque : {task_risk}")
        run_risk.font.color.rgb = _risk_color(task_risk)
        run_risk.bold = True
        meta.add_run(f"  ·  Confiance : {int(float(task_conf)*100)}%" if task_conf else "")

        # Contenu de l'analyse
        if result.get("analysis"):
            _add_heading(doc, "Analyse", level=3)
            _add_paragraph(doc, result["analysis"])

        if result.get("legal_basis"):
            _add_heading(doc, "Base légale", level=3)
            _add_paragraph(doc, result["legal_basis"])

        if result.get("key_points"):
            _add_heading(doc, "Points clés", level=3)
            for point in result["key_points"]:
                p   = doc.add_paragraph(style="List Bullet")
                run = p.add_run(point)
                run.font.size = Pt(11)

        if result.get("recommendation"):
            _add_heading(doc, "Recommandation", level=3)
            _add_paragraph(doc, result["recommendation"])

        if result.get("expert_note"):
            p   = doc.add_paragraph()
            run = p.add_run(f"Note experte : {result['expert_note']}")
            run.font.italic    = True
            run.font.color.rgb = COLOR_MUTED
            run.font.size      = Pt(10)

        doc.add_paragraph()

    # ── Synthèse globale ──────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "3. Synthèse et recommandations globales", level=1)

    if exec_summary:
        _add_paragraph(doc, exec_summary)
    doc.add_paragraph()

    # Tableau de synthèse des risques par agent
    done_tasks = [t for t in tasks if t.get("status") == "done"]
    if done_tasks:
        _add_heading(doc, "Matrice de risque par agent", level=2)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        headers   = ["Agent", "Tâche", "Risque", "Confiance"]
        hdr       = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            _set_cell_bg(hdr[i], "1E293B")
            run = hdr[i].paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

        for task in done_tasks:
            rc = tbl.add_row().cells
            rc[0].text = (task.get("agent_type") or "").upper()
            rc[1].text = (task.get("task_title") or "")[:50]
            rc[2].text = task.get("risk_level") or "—"
            conf = task.get("confidence", 0)
            rc[3].text = f"{int(float(conf)*100)}%" if conf else "—"

    doc.add_paragraph()
    _add_paragraph(
        doc,
        f"Rapport généré automatiquement par EnterpriseCore AI — {datetime.now().strftime('%d/%m/%Y %H:%M')}",
        muted=True
    )

    # ── Export bytes ──────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
